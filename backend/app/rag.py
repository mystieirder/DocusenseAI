"""
RAG pipeline — the retrieval half of DocuSense.

Covers FR-01 (ingest → extract → chunk → embed) and FR-03 (hybrid retrieval):

  extract_pages()   multi-format text extraction with normalized bounding boxes
                    (PyMuPDF blocks; density-based OCR fallback for scanned pages;
                    docx, images, plain text).
  chunk_pages()     semantic-ish chunking (~1000 chars, 150 overlap) carrying a
                    union bbox per chunk so citations can be drawn on the page.
  embed_texts()     dense embeddings via sentence-transformers all-MiniLM-L6-v2 (384-d).
  process_document() the whole ingest path, returning (page_count, chunks[]).
  retrieve()        hybrid search: pgvector cosine (dense) + Postgres full-text
                    ts_rank (sparse), fused with weighted Reciprocal Rank Fusion.
  build_citations() collapse retrieved chunks into per-page citation boxes.

Heavy deps (fitz, pytesseract, sentence_transformers) are imported lazily so the
module imports cleanly even where they're absent.
"""
from __future__ import annotations

import io
import logging
import threading
from typing import List, Optional, Tuple

from sqlalchemy import func
from sqlalchemy.orm import Session

from .config import settings
from .models import DocumentChunk

log = logging.getLogger("docusense.rag")

# ── Embedding model (lazy singleton) ─────────────────────────────────────
_model = None
_model_lock = threading.Lock()


def get_embedder():
    """Load the local sentence-transformers model once (EMBED_BACKEND=local)."""
    global _model
    if _model is None:
        with _model_lock:
            if _model is None:
                from sentence_transformers import SentenceTransformer
                log.info("Loading local embedding model %s", settings.EMBED_MODEL)
                _model = SentenceTransformer(settings.EMBED_MODEL)
    return _model


def _l2_normalize(rows) -> List[list]:
    """L2-normalize a batch of vectors (list[list[float]]) → list[list[float]]."""
    import numpy as np
    arr = np.asarray(rows, dtype="float32")
    if arr.ndim != 2 or arr.size == 0:
        return [list(map(float, r)) for r in rows]
    norms = np.linalg.norm(arr, axis=1, keepdims=True)
    norms[norms == 0] = 1.0
    return (arr / norms).astype("float32").tolist()


def embed_texts(texts, *, is_query: bool = False) -> Optional[List[list]]:
    """
    Embed texts into dense vectors using the configured backend.

    Returns list[list[float]] (L2-normalized), or None when EMBED_BACKEND="keyword"
    (dense disabled) or there is nothing to embed. Returning None — rather than
    raising — lets the ingest path still store chunks (sparse retrieval keeps working).
    """
    if not texts:
        return None
    backend = settings.EMBED_BACKEND.lower()

    if backend == "keyword":
        return None

    if backend == "gemini":
        from . import llm
        task = "RETRIEVAL_QUERY" if is_query else "RETRIEVAL_DOCUMENT"
        vecs = llm.embed(list(texts), dim=settings.GEMINI_EMBED_DIM, task_type=task)
        return _l2_normalize(vecs)

    if backend == "mistral":
        from . import llm
        # mistral-embed has no query/document task distinction; one endpoint covers both.
        vecs = llm.embed_mistral(list(texts))
        return _l2_normalize(vecs)

    # default: local sentence-transformers (encoder already L2-normalizes)
    import numpy as np
    model = get_embedder()
    vecs = model.encode(list(texts), normalize_embeddings=True, convert_to_numpy=True,
                        batch_size=32, show_progress_bar=False)
    return np.asarray(vecs, dtype="float32").tolist()


def embed_query(text: str) -> Optional[list]:
    """Embed a single query string → python list (pgvector-friendly), or None."""
    vecs = embed_texts([text], is_query=True)
    return vecs[0] if vecs else None


# ── Extraction ───────────────────────────────────────────────────────────
def _printable_density(text: str) -> float:
    """Fraction of non-space printable characters — low ⇒ likely a scanned page."""
    if not text:
        return 0.0
    printable = sum(1 for ch in text if ch.isprintable() and not ch.isspace())
    return printable / max(len(text), 1)


def _tesseract_ocr(img) -> str:
    try:
        import pytesseract
        return pytesseract.image_to_string(img).strip()
    except Exception as e:  # tesseract missing or failed — degrade gracefully
        log.warning("Tesseract OCR failed: %s", e)
        return ""


def _page_to_jpeg(img, maxdim: int = 2000, quality: int = 85) -> bytes:
    """Encode a PIL page image to JPEG bytes, downscaling oversized scans.

    A 150-DPI A4 page is already ~1240×1754, so the cap only shrinks very large
    uploaded images — keeping the cloud OCR request payload/latency small without
    hurting transcription quality.
    """
    im = img if img.mode == "RGB" else img.convert("RGB")
    w, h = im.size
    if max(w, h) > maxdim:
        scale = maxdim / float(max(w, h))
        im = im.resize((max(1, int(w * scale)), max(1, int(h * scale))))
    buf = io.BytesIO()
    im.save(buf, format="JPEG", quality=quality)
    return buf.getvalue()


def _jpeg_to_pil(data: bytes):
    from PIL import Image
    return Image.open(io.BytesIO(data))


def _is_quota_error(e: Exception) -> bool:
    s = str(e).lower()
    return any(k in s for k in ("quota", "rate limit", "429", "exhausted", "resource_exhausted"))


def _cloud_ocr_available() -> bool:
    """Whether the active LLM provider can run cloud vision OCR (its API key is set)."""
    provider = (settings.LLM_PROVIDER or "gemini").strip().lower()
    if provider == "mistral":
        return bool(settings.MISTRAL_API_KEY)
    return bool(settings.GEMINI_API_KEY)


def _cloud_ocr(img) -> str:
    """Transcribe a single PIL image via the active provider's multimodal model."""
    from . import llm
    return llm.ocr_image(_page_to_jpeg(img), mime_type="image/jpeg")


def _ocr_image(img) -> str:
    """
    OCR one page/image, choosing the backend per settings.OCR_BACKEND:
      "gemini"/"mistral" → cloud vision via the active LLM provider (needs its key)
      "tesseract"        → local Tesseract binary
      "auto"             → cloud vision when the active provider's key is set, else Tesseract
    A cloud failure (quota / network) degrades to Tesseract, then to "".
    """
    backend = (settings.OCR_BACKEND or "auto").lower()
    use_cloud = backend != "tesseract" and _cloud_ocr_available()
    if use_cloud:
        try:
            return _cloud_ocr(img).strip()
        except Exception as e:
            log.warning("Cloud OCR failed (%s); falling back to Tesseract", e)
    return _tesseract_ocr(img)


def _ocr_pages(images: List[bytes]) -> List[str]:
    """
    OCR a list of page images (JPEG bytes) → one transcription per image.

    Honors settings.OCR_BACKEND. For a cloud backend, pages are transcribed in
    batches of OCR_BATCH_PAGES per request — this is what keeps a scanned PDF under
    a free-tier request limit (e.g. Gemini's ~20 req/min). Resilience, in order:
      • a batch that errors or comes back malformed is retried page-by-page;
      • a per-page cloud failure degrades to local Tesseract, then to "";
      • once a hard quota/rate error is seen, we stop calling the cloud for the rest
        of this document and use Tesseract directly (no point burning doomed requests).
    """
    n = len(images)
    results: List[str] = [""] * n
    if n == 0:
        return results

    backend = (settings.OCR_BACKEND or "auto").lower()
    use_cloud = backend != "tesseract" and _cloud_ocr_available()
    if not use_cloud:
        for i, data in enumerate(images):
            results[i] = _tesseract_ocr(_jpeg_to_pil(data))
        return results

    from . import llm
    batch = max(1, settings.OCR_BATCH_PAGES)
    cloud_dead = False
    i = 0
    while i < n:
        group = images[i:i + batch]
        if cloud_dead:
            for j, data in enumerate(group):
                results[i + j] = _tesseract_ocr(_jpeg_to_pil(data))
            i += batch
            continue
        try:
            texts = llm.ocr_images(group)
            for j, t in enumerate(texts):
                results[i + j] = (t or "").strip()
        except Exception as e:
            # Batch failed (quota, network, or a malformed multi-page split). Retry
            # each page on its own so one bad page can't lose the whole batch.
            if _is_quota_error(e):
                log.warning("Cloud OCR hit a quota/rate wall (%s); switching this doc to Tesseract", e)
                cloud_dead = True
                for j, data in enumerate(group):
                    results[i + j] = _tesseract_ocr(_jpeg_to_pil(data))
                i += batch
                continue
            log.warning("Batch OCR failed (%s); retrying %d page(s) individually", e, len(group))
            for j, data in enumerate(group):
                try:
                    results[i + j] = llm.ocr_image(data).strip()
                except Exception as e2:
                    if _is_quota_error(e2):
                        cloud_dead = True
                    else:
                        log.warning("Cloud OCR failed on a page (%s); using Tesseract", e2)
                    results[i + j] = _tesseract_ocr(_jpeg_to_pil(data))
        i += batch
    return results


def _extract_pdf(content: bytes) -> Tuple[List[dict], int]:
    import fitz  # PyMuPDF
    from PIL import Image

    doc = fitz.open(stream=content, filetype="pdf")
    try:
        page_count = doc.page_count
        pages: List[dict] = [{"page": i + 1, "blocks": []} for i in range(page_count)]
        ocr_imgs: List[bytes] = []     # JPEG bytes for pages that need OCR
        ocr_targets: List[int] = []    # parallel page indices for ocr_imgs

        for i in range(page_count):
            page = doc.load_page(i)
            width = page.rect.width or 1.0
            height = page.rect.height or 1.0

            raw_text = page.get_text().strip()
            needs_ocr = (not raw_text) or (_printable_density(raw_text) < settings.LOW_TEXT_DENSITY)

            if needs_ocr and i < settings.OCR_PAGE_LIMIT:
                try:
                    pix = page.get_pixmap(dpi=150)
                    img = Image.open(io.BytesIO(pix.tobytes("png")))
                    ocr_imgs.append(_page_to_jpeg(img))
                    ocr_targets.append(i)
                    continue  # text is filled in by the batched OCR pass below
                except Exception as e:
                    log.warning("Rasterization failed on page %d: %s", i + 1, e)

            # Digital text page (or OCR not applicable): keep the positioned blocks.
            for b in page.get_text("blocks"):
                # (x0, y0, x1, y1, text, block_no, block_type)
                x0, y0, x1, y1, txt = b[0], b[1], b[2], b[3], b[4]
                txt = (txt or "").strip()
                if not txt:
                    continue
                pages[i]["blocks"].append({
                    "text": txt,
                    "bbox": [x0 / width, y0 / height, x1 / width, y1 / height],
                })

        # OCR pass — batched so a many-page scan stays within the free-tier limit.
        if ocr_imgs:
            texts = _ocr_pages(ocr_imgs)
            for idx, text in zip(ocr_targets, texts):
                text = (text or "").strip()
                if text:
                    pages[idx]["blocks"] = [{"text": text, "bbox": [0.0, 0.0, 1.0, 1.0]}]

        return pages, page_count
    finally:
        doc.close()


def _extract_docx(content: bytes) -> Tuple[List[dict], int]:
    import docx
    d = docx.Document(io.BytesIO(content))
    paragraphs = [p.text for p in d.paragraphs if p.text and p.text.strip()]
    for table in d.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    full = "\n".join(paragraphs).strip()
    blocks = [{"text": full, "bbox": [0.0, 0.0, 1.0, 1.0]}] if full else []
    return [{"page": 1, "blocks": blocks}], 1


def _extract_image(content: bytes) -> Tuple[List[dict], int]:
    from PIL import Image
    try:
        img = Image.open(io.BytesIO(content))
        text = _ocr_image(img)
    except Exception as e:
        log.warning("Image open failed: %s", e)
        text = ""
    blocks = [{"text": text, "bbox": [0.0, 0.0, 1.0, 1.0]}] if text else []
    return [{"page": 1, "blocks": blocks}], 1


def _extract_txt(content: bytes) -> Tuple[List[dict], int]:
    text = content.decode("utf-8", errors="ignore").strip()
    blocks = [{"text": text, "bbox": [0.0, 0.0, 1.0, 1.0]}] if text else []
    return [{"page": 1, "blocks": blocks}], 1


def extract_pages(content: bytes, filename: str) -> Tuple[List[dict], int]:
    """
    Dispatch on file extension → (pages, page_count).
    pages: [{"page": int, "blocks": [{"text": str, "bbox": [x0,y0,x1,y1]}]}]
    """
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        return _extract_pdf(content)
    if name.endswith(".docx"):
        return _extract_docx(content)
    if name.endswith((".png", ".jpg", ".jpeg")):
        return _extract_image(content)
    return _extract_txt(content)


# ── Chunking ──────────────────────────────────────────────────────────────
def _union_bbox(boxes: List[Optional[list]]) -> Optional[list]:
    valid = [b for b in boxes if b]
    if not valid:
        return None
    return [
        min(b[0] for b in valid),
        min(b[1] for b in valid),
        max(b[2] for b in valid),
        max(b[3] for b in valid),
    ]


def chunk_pages(pages: List[dict]) -> List[dict]:
    """
    Group page blocks into ~CHUNK_CHARS-sized chunks (overlap CHUNK_OVERLAP),
    keeping a union bbox so citations can be highlighted. Chunks never span pages.
    Returns [{"page_num","chunk_index","content","bbox"}] with a global chunk_index.
    """
    ch, ov = settings.CHUNK_CHARS, settings.CHUNK_OVERLAP
    chunks: List[dict] = []
    idx = 0

    for pg in pages:
        page_num = pg["page"]
        buf: List[str] = []
        buf_boxes: List[Optional[list]] = []
        cur_len = 0

        def flush():
            nonlocal idx, buf, buf_boxes, cur_len
            if not buf:
                return
            text = "\n".join(buf).strip()
            if text:
                chunks.append({
                    "page_num": page_num,
                    "chunk_index": idx,
                    "content": text,
                    "bbox": _union_bbox(buf_boxes),
                })
                idx += 1
            buf, buf_boxes, cur_len = [], [], 0

        for blk in pg.get("blocks", []):
            t = (blk.get("text") or "").strip()
            if not t:
                continue
            bbox = blk.get("bbox")

            # A single block longer than the window: hard-split with overlap.
            if len(t) > ch:
                flush()
                start = 0
                while start < len(t):
                    piece = t[start:start + ch]
                    chunks.append({
                        "page_num": page_num,
                        "chunk_index": idx,
                        "content": piece,
                        "bbox": bbox,
                    })
                    idx += 1
                    start += max(ch - ov, 1)
                continue

            if cur_len + len(t) > ch and buf:
                flush()
            buf.append(t)
            buf_boxes.append(bbox)
            cur_len += len(t)

        flush()

    return chunks


# ── Full ingest path ────────────────────────────────────────────────────
def process_document(content: bytes, filename: str) -> Tuple[int, List[dict]]:
    """
    Extract → chunk → embed. Returns (page_count, chunks) where each chunk is
    {"page_num","chunk_index","content","bbox","embedding"(list|None)}.
    Embedding failures degrade to sparse-only (chunks still stored).
    """
    pages, page_count = extract_pages(content, filename)
    chunks = chunk_pages(pages)
    if chunks:
        try:
            embeddings = embed_texts([c["content"] for c in chunks])
        except Exception as e:
            log.error("Embedding failed, storing chunks without vectors: %s", e)
            embeddings = None
        for i, c in enumerate(chunks):
            c["embedding"] = embeddings[i] if embeddings else None
    return page_count, chunks


# ── Hybrid retrieval (FR-03) ──────────────────────────────────────────────
def _chunk_dict(c: DocumentChunk, score: float) -> dict:
    return {
        "id": c.id,
        "page_num": c.page_num,
        "chunk_index": c.chunk_index,
        "content": c.content,
        "bbox": c.bbox_json,
        "score": score,
    }


def retrieve(db: Session, doc_id: str, query: str, top_k: Optional[int] = None) -> List[dict]:
    """
    Hybrid dense + sparse retrieval for one document.

    Dense: pgvector cosine distance on the query embedding.
    Sparse: Postgres full-text ts_rank via plainto_tsquery.
    Fusion: weighted Reciprocal Rank Fusion (HYBRID_ALPHA controls dense weight).
    Falls back to leading chunks if neither arm returns anything.
    """
    top_k = top_k or settings.RETRIEVAL_TOP_K
    fetch = max(top_k * 3, 20)

    # Dense arm
    dense_ids: List[int] = []
    qvec = None
    try:
        qvec = embed_query(query)
    except Exception as e:
        log.warning("Query embedding failed, using sparse only: %s", e)
    if qvec is not None:
        rows = (
            db.query(DocumentChunk.id)
            .filter(DocumentChunk.doc_id == doc_id, DocumentChunk.embedding.isnot(None))
            .order_by(DocumentChunk.embedding.cosine_distance(qvec))
            .limit(fetch)
            .all()
        )
        dense_ids = [r[0] for r in rows]

    # Sparse arm
    tsquery = func.plainto_tsquery("english", query)
    srows = (
        db.query(DocumentChunk.id)
        .filter(DocumentChunk.doc_id == doc_id, DocumentChunk.tsv.op("@@")(tsquery))
        .order_by(func.ts_rank(DocumentChunk.tsv, tsquery).desc())
        .limit(fetch)
        .all()
    )
    sparse_ids = [r[0] for r in srows]

    # Weighted RRF fusion
    alpha = settings.HYBRID_ALPHA
    k = 60
    scores: dict = {}
    for rank, cid in enumerate(dense_ids):
        scores[cid] = scores.get(cid, 0.0) + alpha * (1.0 / (k + rank + 1))
    for rank, cid in enumerate(sparse_ids):
        scores[cid] = scores.get(cid, 0.0) + (1.0 - alpha) * (1.0 / (k + rank + 1))

    if not scores:
        rows = (
            db.query(DocumentChunk)
            .filter(DocumentChunk.doc_id == doc_id)
            .order_by(DocumentChunk.page_num, DocumentChunk.chunk_index)
            .limit(top_k)
            .all()
        )
        return [_chunk_dict(c, 0.0) for c in rows]

    top_ids = sorted(scores, key=lambda cid: scores[cid], reverse=True)[:top_k]
    fetched = db.query(DocumentChunk).filter(DocumentChunk.id.in_(top_ids)).all()
    by_id = {c.id: c for c in fetched}
    results = [_chunk_dict(by_id[cid], scores[cid]) for cid in top_ids if cid in by_id]
    # present in reading order
    results.sort(key=lambda r: (r["page_num"], r["chunk_index"]))
    return results


def build_citations(chunks: List[dict]) -> List[dict]:
    """
    Collapse retrieved chunks into per-page citation boxes:
    [{"page": int, "boxes": [[x0,y0,x1,y1], ...]}], ordered by page.
    """
    by_page: dict = {}
    for c in chunks:
        page = c.get("page_num")
        if page is None:
            continue
        box = c.get("bbox")
        by_page.setdefault(page, [])
        if box:
            by_page[page].append([round(float(v), 5) for v in box])
    return [{"page": p, "boxes": by_page[p]} for p in sorted(by_page)]